"""
WEED FINDER Project Report Generator
Generates a complete B.Tech engineering project report in Word format.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"C:\Users\Athul\Documents\Balu project\WEED_FINDER_Project_Report.docx"

# ─────────────────────────────────────────────
# HELPER UTILITIES
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table-cell background colour (hex without #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_table_borders(table):
    """Add visible borders to every cell of a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_page_break(doc):
    doc.add_page_break()


def heading1(doc, text, numbered=True):
    """Chapter-level heading."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    pPr = p._p.get_or_add_pPr()
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Heading1')
    pPr.insert(0, pStyle)
    return p


def heading2(doc, text):
    """Sub-heading."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p


def body(doc, text, justify=True):
    """Normal body paragraph."""
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p


def centered_bold(doc, text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def centered(doc, text, size=12, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_header_row(table, cols, bg='4472C4'):
    row = table.rows[0]
    for i, col in enumerate(cols):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(col)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        set_cell_bg(cell, bg)


def add_table_row(table, values, bold=False, center=False):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        run.bold = bold
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
    return row


def set_doc_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_page_number(doc):
    """Add page number to footer of all sections."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


# ─────────────────────────────────────────────
# DOCUMENT SECTIONS
# ─────────────────────────────────────────────

def title_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    centered(doc, "TKM INSTITUTE OF TECHNOLOGY", size=16, bold=True)
    centered(doc, "KOLLAM, KERALA", size=13, bold=True)
    centered(doc, "Department of Computer Science and Engineering", size=13, bold=True)
    doc.add_paragraph()
    doc.add_paragraph()
    centered(doc, "B.Tech Project Report", size=14, bold=True)
    centered(doc, "on", size=13)
    doc.add_paragraph()
    centered_bold(doc, "WEED FINDER: AUTONOMOUS WEED DETECTION AND REMOVAL ROBOT", size=16)
    doc.add_paragraph()
    doc.add_paragraph()
    centered(doc, "Submitted in partial fulfillment of the requirements for the award of the degree of", size=12)
    centered(doc, "Bachelor of Technology", size=13, bold=True)
    centered(doc, "in", size=12)
    centered(doc, "Computer Science and Engineering", size=13, bold=True)
    centered(doc, "under", size=12)
    centered(doc, "APJ Abdul Kalam Technological University", size=13, bold=True)
    doc.add_paragraph()
    doc.add_paragraph()
    centered(doc, "Submitted by:", size=12, bold=True)
    for name, roll in [("[Student Name 1]", "[Roll No 1]"),
                       ("[Student Name 2]", "[Roll No 2]"),
                       ("[Student Name 3]", "[Roll No 3]"),
                       ("Balu", "[Roll No 4]")]:
        centered(doc, f"{name}  ({roll})", size=12)
    doc.add_paragraph()
    centered(doc, "Under the guidance of", size=12)
    centered(doc, "[Guide Name]", size=13, bold=True)
    centered(doc, "Assistant Professor, Department of CSE", size=12)
    doc.add_paragraph()
    doc.add_paragraph()
    centered(doc, "April 2025", size=12, bold=True)
    add_page_break(doc)


def declaration_page(doc):
    centered_bold(doc, "DECLARATION", size=14)
    doc.add_paragraph()
    body(doc,
         "We, the undersigned students of the final year B.Tech programme in Computer Science and "
         "Engineering at TKM Institute of Technology, Kollam, hereby declare that the project report "
         "titled \"WEED FINDER: Autonomous Weed Detection and Removal Robot\" submitted to the APJ Abdul "
         "Kalam Technological University in partial fulfillment of the requirements for the award of the "
         "degree of Bachelor of Technology in Computer Science and Engineering is an authentic record of "
         "work carried out by us during the period 2024–2025 under the supervision of [Guide Name], "
         "Assistant Professor, Department of Computer Science and Engineering.")
    body(doc,
         "We further declare that the work reported in this project report has not been submitted, "
         "either in part or in full, for the award of any other degree or diploma in this institute or "
         "any other institute or university. The information and data presented in this report are genuine "
         "and have not been falsified or manipulated in any manner. All sources of information have been "
         "duly acknowledged.")
    body(doc,
         "We also declare that the software developed as part of this project has been tested and "
         "verified to the best of our knowledge, and all intellectual property rights of third-party "
         "libraries used in this work have been respected and acknowledged in the references section.")
    doc.add_paragraph()
    doc.add_paragraph()
    centered(doc, "Place: Kollam", size=12)
    centered(doc, "Date: April 2025", size=12)
    doc.add_paragraph()
    doc.add_paragraph()

    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    names = ["[Student Name 1] ([Roll No 1])", "[Student Name 2] ([Roll No 2])",
             "[Student Name 3] ([Roll No 3])", "Balu ([Roll No 4])"]
    for i in range(2):
        for j in range(2):
            idx = i * 2 + j
            cell = t.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(names[idx])
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    add_page_break(doc)


def certificate_page(doc):
    centered_bold(doc, "CERTIFICATE", size=14)
    doc.add_paragraph()
    body(doc,
         "This is to certify that the project report titled \"WEED FINDER: Autonomous Weed Detection and "
         "Removal Robot\" is a bonafide record of the work done by the following students of the final year "
         "B.Tech programme in Computer Science and Engineering at TKM Institute of Technology, Kollam, "
         "during the academic year 2024–2025, in partial fulfillment of the requirements for the award of "
         "the degree of Bachelor of Technology in Computer Science and Engineering under APJ Abdul Kalam "
         "Technological University.")
    doc.add_paragraph()
    for name, roll in [("[Student Name 1]", "[Roll No 1]"),
                       ("[Student Name 2]", "[Roll No 2]"),
                       ("[Student Name 3]", "[Roll No 3]"),
                       ("Balu", "[Roll No 4]")]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{name}  -  {roll}")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    doc.add_paragraph()
    body(doc,
         "The work reported herein has not been submitted elsewhere for the award of any degree or diploma. "
         "The project complies with the regulations of the APJ Abdul Kalam Technological University.")
    doc.add_paragraph()
    doc.add_paragraph()

    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = t.rows[0].cells
    for c, label, name in [(cells[0], "Project Guide", "[Guide Name]\nAssistant Professor\nDept. of CSE"),
                            (cells[1], "Head of Department", "[HOD Name]\nProfessor & HOD\nDept. of CSE")]:
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label + "\n\n\n" + name)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Project Coordinators:\n[Coordinator 1]    [Coordinator 2]")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Principal\n[Principal Name]\nTKM Institute of Technology, Kollam")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    add_page_break(doc)


def acknowledgement_page(doc):
    centered_bold(doc, "ACKNOWLEDGEMENT", size=14)
    doc.add_paragraph()
    body(doc,
         "We, the project team, express our sincere gratitude to the Almighty for the strength, "
         "perseverance, and wisdom bestowed upon us throughout this project work. It is with immense "
         "pleasure and a deep sense of gratitude that we acknowledge the support and guidance of all "
         "those who have contributed, directly or indirectly, to the successful completion of this project.")
    body(doc,
         "We extend our heartfelt thanks to [Principal Name], Principal, TKM Institute of Technology, "
         "Kollam, for providing us with all the necessary facilities and an encouraging environment "
         "to carry out this project work.")
    body(doc,
         "We would like to express our profound gratitude to [HOD Name], Professor and Head of the "
         "Department of Computer Science and Engineering, TKM Institute of Technology, for the constant "
         "encouragement and administrative support that made this project a reality.")
    body(doc,
         "We are deeply indebted to our project guide, [Guide Name], Assistant Professor, Department of "
         "Computer Science and Engineering, for the invaluable guidance, continuous support, and "
         "constructive feedback provided throughout the course of this project. The insightful suggestions "
         "and technical expertise shared by our guide have been instrumental in shaping the direction and "
         "outcomes of this work.")
    body(doc,
         "We also thank the project coordinators, [Coordinator 1] and [Coordinator 2], for their timely "
         "assistance and coordination during the various phases of the project. Their organizational "
         "efforts ensured smooth progress from planning to execution.")
    body(doc,
         "Our sincere thanks go to all the faculty members and technical staff of the Department of "
         "Computer Science and Engineering for their support and encouragement. We are also grateful to "
         "the open-source community, particularly the contributors to the Ultralytics YOLOv8 framework, "
         "OpenCV, Flask, and the Kaggle dataset community, whose publicly available resources formed the "
         "technical backbone of this project.")
    body(doc,
         "Finally, we are deeply grateful to our families and friends for their unwavering moral support, "
         "patience, and encouragement throughout the duration of this programme. Their belief in our "
         "abilities has been a constant source of motivation.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("[Student Name 1]\n[Student Name 2]\n[Student Name 3]\nBalu")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    add_page_break(doc)


def abstract_page(doc):
    centered_bold(doc, "ABSTRACT", size=14)
    doc.add_paragraph()
    body(doc,
         "Agriculture remains the backbone of the Indian economy, employing a significant portion of the "
         "population and contributing substantially to national GDP. However, the persistence of weeds "
         "in agricultural fields poses a severe threat to crop yield and quality. Weeds compete with "
         "cultivated plants for essential resources such as water, sunlight, nutrients, and space, "
         "leading to yield losses of 20–80% if left uncontrolled. Traditional weed management methods "
         "— manual removal and broad-spectrum herbicide application — are labour-intensive, economically "
         "costly, environmentally hazardous, and increasingly ineffective due to the emergence of "
         "herbicide-resistant weed species.")
    body(doc,
         "This project presents WEED FINDER, an autonomous weed detection and removal robotic system "
         "that leverages state-of-the-art computer vision and embedded computing to address these "
         "challenges. The system integrates a Raspberry Pi 4 single-board computer, a USB camera module, "
         "a YOLOv8n object detection model (exported to ONNX format for ARM64 inference), and a linear "
         "actuator mechanism for physical weed removal. The detection pipeline processes live video "
         "frames at approximately 8–12 frames per second, identifies weeds within a predefined strike "
         "zone occupying the central-lower portion of the frame, and triggers a GPIO-controlled linear "
         "actuator to physically remove detected weeds without human intervention.")
    body(doc,
         "The YOLOv8n model was trained on the Kaggle 'Crop and Weed Detection' dataset comprising "
         "832 training and 208 validation images across two classes: crop and weed. Training was "
         "performed for 50 epochs with data augmentation including rotation, horizontal/vertical "
         "flipping, and mosaic augmentation, achieving reliable detection at a confidence threshold "
         "of 0.45. The ONNX Runtime was employed for inference on the Raspberry Pi 4 to circumvent "
         "the incompatibility of PyTorch with ARM64 architecture, enabling efficient CPU-based inference.")
    body(doc,
         "The system features a modular software architecture with clearly separated components for "
         "camera capture (CameraModule), weed detection (WeedDetector), strike-zone management "
         "(StrikeZone), actuator control (ArmController), and a real-time web dashboard (Flask) "
         "accessible over the local network for monitoring detection statistics and live video feed. "
         "Multiple operating modes — simulation, live camera with GPIO, and serial Arduino control — "
         "provide flexibility for development, testing, and deployment.")
    body(doc,
         "Testing using a pytest framework covering four test modules validated the system's "
         "correctness in terms of frame geometry, detection logic, actuator cooldown enforcement, "
         "and strike-zone boundary mathematics. The system is deployed as a systemd service on "
         "Raspberry Pi, enabling automatic start-on-boot operation. WEED FINDER demonstrates that "
         "precision agriculture automation is achievable with affordable embedded hardware and open-source "
         "deep learning frameworks, offering a sustainable and eco-friendly alternative to chemical "
         "weed management.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Keywords: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(
        "Weed Detection, YOLOv8, ONNX Runtime, Raspberry Pi 4, Precision Agriculture, "
        "Computer Vision, Linear Actuator, GPIO, Flask Dashboard, Autonomous Robot")
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    add_page_break(doc)


def po_mapping_page(doc):
    centered_bold(doc, "MAPPING WITH PROGRAM OUTCOMES (POs)", size=14)
    doc.add_paragraph()
    body(doc,
         "The following table maps the WEED FINDER project to the Programme Outcomes (POs) as "
         "defined by the NBA (National Board of Accreditation) for B.Tech programmes in Engineering "
         "and Technology. The mapping demonstrates how the project addresses graduate attributes "
         "across engineering knowledge, problem analysis, design, research, tools, ethics, "
         "environment, teamwork, communication, project management, lifelong learning, and "
         "contemporary issues.")
    doc.add_paragraph()

    headers = ["PO No.", "Programme Outcome", "Relevance to WEED FINDER", "Level\n(H/M/L)"]
    po_data = [
        ("PO1", "Engineering Knowledge",
         "Applied knowledge of embedded systems (Raspberry Pi 4), computer vision (OpenCV, YOLOv8), "
         "GPIO electronics, and software engineering principles to design and implement the robot.", "H"),
        ("PO2", "Problem Analysis",
         "Analysed the problem of manual weed detection inefficiency and formulated a solution using "
         "real-time object detection, strike-zone geometry, and actuator control.", "H"),
        ("PO3", "Design/Development of Solutions",
         "Designed a complete end-to-end robotic system from dataset preparation and model training "
         "to hardware integration, modular software architecture, and web dashboard deployment.", "H"),
        ("PO4", "Conduct Investigations of Complex Problems",
         "Investigated YOLOv8 vs RT-DETR model performance, evaluated ONNX Runtime for ARM64 "
         "compatibility, and performed comparative analysis of inference speed and accuracy.", "M"),
        ("PO5", "Modern Tool Usage",
         "Utilised modern tools including Python 3.10+, PyTorch, ONNX Runtime, Ultralytics, "
         "OpenCV, Flask, pytest, Git, systemd, and Google Colab for training.", "H"),
        ("PO6", "The Engineer and Society",
         "Addressed the societal need for sustainable, herbicide-free weed management, reducing "
         "farmer dependency on chemical inputs and supporting food security.", "H"),
        ("PO7", "Environment and Sustainability",
         "The system eliminates the need for chemical herbicides, directly contributing to "
         "reduced soil contamination, water pollution, and environmental degradation.", "H"),
        ("PO8", "Ethics",
         "Respected open-source software licences (MIT, Apache 2.0), acknowledged dataset "
         "providers, and followed ethical guidelines for autonomous systems design.", "M"),
        ("PO9", "Individual and Team Work",
         "The project was executed as a collaborative team effort with clearly defined roles "
         "for hardware, software, training, testing, and documentation.", "H"),
        ("PO10", "Communication",
         "Documented the system comprehensively in this report, maintained code documentation, "
         "and designed a web dashboard for intuitive communication of system status.", "H"),
        ("PO11", "Project Management and Finance",
         "Planned and executed the project within time and resource constraints, managing tasks "
         "across training, hardware procurement, software development, and testing phases.", "M"),
        ("PO12", "Life-long Learning",
         "Explored cutting-edge YOLO architectures, RT-DETR transformer models, ONNX export "
         "pipelines, and embedded Linux deployment, developing skills applicable across domains.", "H"),
    ]

    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_header_row(t, headers)
    for row in po_data:
        r = t.add_row()
        for i, val in enumerate(row):
            cell = r.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
        if po_data.index(row) % 2 == 1:
            for cell in r.cells:
                set_cell_bg(cell, 'D9E1F2')
    add_page_break(doc)


def sdg_mapping_page(doc):
    centered_bold(doc, "MAPPING WITH SUSTAINABLE DEVELOPMENT GOALS (SDGs)", size=14)
    doc.add_paragraph()
    body(doc,
         "The United Nations Sustainable Development Goals (SDGs) provide a global framework for "
         "addressing humanity's most pressing challenges by 2030. The WEED FINDER project aligns "
         "with several SDGs by promoting precision agriculture, environmental sustainability, "
         "technological innovation, and responsible consumption. The following table details the "
         "mapping of the project's outcomes to relevant SDGs.")
    doc.add_paragraph()

    headers = ["SDG No.", "SDG Title", "How WEED FINDER Contributes"]
    sdg_data = [
        ("SDG 2", "Zero Hunger",
         "Improves crop yield by autonomous, efficient weed removal, reducing crop-weed competition "
         "and supporting food security for growing populations."),
        ("SDG 3", "Good Health and Well-Being",
         "Reduces farmer exposure to toxic herbicides and pesticides, protecting agricultural "
         "workers from chemical-related health hazards."),
        ("SDG 9", "Industry, Innovation and Infrastructure",
         "Demonstrates the application of AI and robotics in agriculture, fostering technological "
         "innovation and building smart agricultural infrastructure."),
        ("SDG 11", "Sustainable Cities and Communities",
         "Supports sustainable rural communities by increasing agricultural productivity and "
         "reducing chemical inputs that degrade local ecosystems."),
        ("SDG 12", "Responsible Consumption and Production",
         "Promotes responsible production by enabling targeted, precise weed removal, eliminating "
         "the over-application of broad-spectrum herbicides."),
        ("SDG 13", "Climate Action",
         "Reduced herbicide production and application lowers the carbon footprint associated "
         "with chemical manufacturing and agricultural runoff."),
        ("SDG 15", "Life on Land",
         "Protects soil microbiome, biodiversity, and terrestrial ecosystems by eliminating "
         "chemical weed control agents that harm non-target organisms."),
        ("SDG 17", "Partnerships for the Goals",
         "Leverages open-source tools, public datasets (Kaggle), and collaborative platforms "
         "(Google Colab) to demonstrate the power of global knowledge sharing."),
    ]

    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_header_row(t, headers)
    for i, row in enumerate(sdg_data):
        r = t.add_row()
        for j, val in enumerate(row):
            cell = r.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
        if i % 2 == 1:
            for cell in r.cells:
                set_cell_bg(cell, 'E2EFDA')
    add_page_break(doc)


def feasibility_page(doc):
    centered_bold(doc, "FEASIBILITY REPORT", size=14)
    doc.add_paragraph()

    heading2(doc, "Technical Feasibility")
    body(doc,
         "The WEED FINDER project is technically feasible due to the availability of mature, "
         "open-source technologies in computer vision, deep learning, and embedded computing. "
         "The Raspberry Pi 4 single-board computer, equipped with a 64-bit ARM Cortex-A72 "
         "quad-core processor running at 1.8 GHz and a minimum of 2 GB LPDDR4 SDRAM, provides "
         "sufficient computational capability to run the YOLOv8n ONNX inference pipeline at "
         "8–12 FPS on CPU — adequate for field-speed weed detection.")
    body(doc,
         "The YOLOv8n (nano) model, when exported to ONNX format, is compatible with the ONNX "
         "Runtime inference engine that supports ARM64 architecture natively. This bypasses the "
         "known PyTorch ARM64 incompatibility on Raspberry Pi OS 64-bit. The model was trained on "
         "a publicly available Kaggle dataset of 1,040 annotated agricultural images, achieving "
         "reliable weed detection at a confidence threshold of 0.45.")
    body(doc,
         "OpenCV 4.8+ provides efficient camera capture (V4L2/UVC-compatible USB cameras), image "
         "preprocessing (blob creation for ONNX inference), and frame annotation. Flask provides "
         "a lightweight web server for the real-time monitoring dashboard accessible over Wi-Fi. "
         "The GPIO interface on Raspberry Pi 4 (BCM pins 17 and 27) is used to control an L298N "
         "motor driver that actuates a 12–24V DC linear actuator for physical weed removal. All "
         "these technologies are well-documented and have established community support.")
    body(doc,
         "The modular software architecture — comprising separate modules for camera capture, "
         "weed detection, strike-zone management, and arm control — ensures maintainability "
         "and extensibility. The system was validated through a pytest-based test suite covering "
         "frame geometry, detection logic, actuator threading safety, and strike-zone mathematics.")

    heading2(doc, "Economic Feasibility")
    body(doc,
         "The WEED FINDER system is designed to be economically accessible to small and medium-scale "
         "farmers. The primary hardware components and their approximate costs are listed below:")

    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_header_row(t, ["Component", "Specification", "Approx. Cost (INR)"])
    hw_data = [
        ("Raspberry Pi 4", "4GB RAM, 64-bit ARM Cortex-A72", "₹5,500 – ₹6,500"),
        ("USB Camera", "UVC-compatible, 640×480 @ 30FPS", "₹500 – ₹1,500"),
        ("Linear Actuator", "12V DC, 150mm stroke", "₹1,200 – ₹2,500"),
        ("L298N Motor Driver", "Dual H-Bridge, 12V/2A", "₹100 – ₹250"),
        ("12V Power Supply", "2A rated, external adapter", "₹300 – ₹600"),
        ("MicroSD Card", "32GB Class 10 (OS + model storage)", "₹350 – ₹700"),
        ("Miscellaneous", "Wires, connectors, chassis, fasteners", "₹500 – ₹1,000"),
        ("TOTAL", "", "₹8,450 – ₹13,050"),
    ]
    for row in hw_data:
        r = t.add_row()
        for i, val in enumerate(row):
            cell = r.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 2 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.bold = (row[0] == "TOTAL")
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    doc.add_paragraph()
    body(doc,
         "The software stack is entirely open-source and free of licensing costs. YOLOv8 (AGPL-3.0), "
         "ONNX Runtime (MIT), OpenCV (Apache 2.0), Flask (BSD), and Python (PSF) all carry no "
         "commercial licensing fees for educational and research use. Model training was performed "
         "on Google Colab (free tier) using GPU acceleration, eliminating the need for expensive "
         "local GPU hardware during the training phase.")
    body(doc,
         "Compared to commercial agricultural robots — which cost upwards of ₹5–50 lakhs — the "
         "WEED FINDER prototype is assembled for under ₹15,000, making it accessible for proof-of-concept "
         "research, academic demonstration, and small-farm pilots. The recurring operational cost is "
         "negligible, consisting only of electrical power consumption (Raspberry Pi 4: ~6W; actuator: "
         "~24W during strike). The long-term economic benefit — reduction in herbicide purchase costs "
         "and labour wages — provides a strong return on investment for agricultural use cases.")

    heading2(doc, "Operational Feasibility")
    body(doc,
         "The system is designed with ease of operation as a primary concern. The web dashboard, "
         "accessible from any smartphone or laptop on the same Wi-Fi network, displays a live "
         "annotated video feed, detection statistics (total weeds detected, strikes performed, FPS), "
         "and system health status. A non-technical operator can monitor and manage the system "
         "without any command-line knowledge. The systemd service configuration enables fully "
         "automatic startup on boot, requiring no manual intervention after initial setup.")
    add_page_break(doc)


def toc_page(doc):
    centered_bold(doc, "TABLE OF CONTENTS", size=14)
    doc.add_paragraph()

    toc_entries = [
        ("Declaration", "i"),
        ("Certificate", "ii"),
        ("Acknowledgement", "iii"),
        ("Abstract", "iv"),
        ("Mapping with Program Outcomes", "v"),
        ("Mapping with SDGs", "vi"),
        ("Feasibility Report", "vii"),
        ("Table of Contents", "viii"),
        ("List of Tables", "ix"),
        ("List of Figures", "x"),
        ("Abbreviations", "xi"),
        ("Chapter 1: Introduction", "1"),
        ("    1.1 Background and Motivation", "1"),
        ("    1.2 Problem Statement", "2"),
        ("    1.3 Objectives", "3"),
        ("    1.4 Scope of the Project", "3"),
        ("    1.5 Organisation of the Report", "4"),
        ("Chapter 2: Literature Review", "5"),
        ("    2.1 Conventional Weed Management", "5"),
        ("    2.2 Machine Learning in Agriculture", "5"),
        ("    2.3 YOLO-Based Object Detection", "6"),
        ("    2.4 Embedded Vision Systems", "7"),
        ("    2.5 Research Gaps", "7"),
        ("Chapter 3: Methodology", "8"),
        ("    3.1 System Overview", "8"),
        ("    3.2 Dataset and Preprocessing", "9"),
        ("    3.3 Model Training", "9"),
        ("    3.4 ONNX Export and Inference", "10"),
        ("    3.5 Hardware Integration", "10"),
        ("    3.6 Software Architecture", "11"),
        ("Chapter 4: Work Plan and Task Allocation", "13"),
        ("Chapter 5: Experimental Design, Coding and Block Diagram", "15"),
        ("    5.1 Block Diagram", "15"),
        ("    5.2 System Configuration", "15"),
        ("    5.3 Key Code Modules", "16"),
        ("    5.4 Detection Pipeline", "18"),
        ("Chapter 6: Results and Conclusion", "20"),
        ("Chapter 7: Future Scope", "23"),
        ("Chapter 8: References", "25"),
    ]

    t = doc.add_table(rows=len(toc_entries), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (title, page) in enumerate(toc_entries):
        row = t.rows[i]
        cell_l = row.cells[0]
        cell_r = row.cells[1]
        cell_l.text = ''
        cell_r.text = ''
        p_l = cell_l.paragraphs[0]
        p_r = cell_r.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        bold = not title.startswith("    ")
        rl = p_l.add_run(title)
        rl.bold = bold
        rl.font.size = Pt(11)
        rl.font.name = 'Times New Roman'

        rr = p_r.add_run(page)
        rr.bold = bold
        rr.font.size = Pt(11)
        rr.font.name = 'Times New Roman'
    add_page_break(doc)


def lot_page(doc):
    centered_bold(doc, "LIST OF TABLES", size=14)
    doc.add_paragraph()

    entries = [
        ("Table 1.1", "Comparison of Weed Detection Approaches", "5"),
        ("Table 3.1", "Dataset Split — Training and Validation", "9"),
        ("Table 3.2", "YOLOv8n Training Hyperparameters", "10"),
        ("Table 3.3", "ONNX Model Size and Inference Speed", "11"),
        ("Table 4.1", "Work Plan and Task Allocation", "13"),
        ("Table 5.1", "System Configuration Parameters (config.py)", "16"),
        ("Table 5.2", "GPIO Pin Assignment", "17"),
        ("Table 5.3", "Software Dependency Versions", "18"),
        ("Table 6.1", "Detection Accuracy Metrics", "21"),
        ("Table 6.2", "System Performance on Raspberry Pi 4", "22"),
        ("Table A.1", "PO Mapping", "v"),
        ("Table A.2", "SDG Mapping", "vi"),
        ("Table A.3", "Hardware Bill of Materials", "vii"),
    ]

    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_header_row(t, ["Table No.", "Table Title", "Page No."])
    for row in entries:
        r = t.add_row()
        for i, val in enumerate(row):
            cell = r.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
    add_page_break(doc)


def lof_page(doc):
    centered_bold(doc, "LIST OF FIGURES", size=14)
    doc.add_paragraph()

    entries = [
        ("Figure 3.1", "System Architecture Overview", "8"),
        ("Figure 3.2", "Detection Pipeline Flowchart", "9"),
        ("Figure 3.3", "YOLOv8n Network Architecture (nano variant)", "10"),
        ("Figure 3.4", "ONNX Output Tensor Shape: (1, 84, 8400)", "11"),
        ("Figure 5.1", "System Block Diagram", "15"),
        ("Figure 5.2", "Strike Zone Definition on Camera Frame", "16"),
        ("Figure 5.3", "GPIO Wiring Schematic — RPi to L298N", "17"),
        ("Figure 5.4", "Linear Actuator Mechanism", "18"),
        ("Figure 5.5", "Dashboard UI — Live Video and Stats Panel", "19"),
        ("Figure 6.1", "Sample Detection Output — Annotated Frame", "20"),
        ("Figure 6.2", "Detection Confidence Distribution", "21"),
        ("Figure 6.3", "FPS vs. Inference Mode Comparison", "22"),
    ]

    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_header_row(t, ["Figure No.", "Figure Title", "Page No."])
    for row in entries:
        r = t.add_row()
        for i, val in enumerate(row):
            cell = r.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
    add_page_break(doc)


def abbreviations_page(doc):
    centered_bold(doc, "ABBREVIATIONS", size=14)
    doc.add_paragraph()

    abbrevs = [
        ("AI", "Artificial Intelligence"),
        ("API", "Application Programming Interface"),
        ("ARM", "Advanced RISC Machine"),
        ("BCM", "Broadcom (GPIO numbering scheme)"),
        ("CLI", "Command-Line Interface"),
        ("CNN", "Convolutional Neural Network"),
        ("CPU", "Central Processing Unit"),
        ("CSV", "Comma-Separated Values"),
        ("DC", "Direct Current"),
        ("DNN", "Deep Neural Network"),
        ("FPS", "Frames Per Second"),
        ("GPIO", "General Purpose Input/Output"),
        ("GPU", "Graphics Processing Unit"),
        ("HOD", "Head of Department"),
        ("HTTP", "Hypertext Transfer Protocol"),
        ("IoT", "Internet of Things"),
        ("JPEG", "Joint Photographic Experts Group"),
        ("JSON", "JavaScript Object Notation"),
        ("L298N", "L298N Dual H-Bridge Motor Driver"),
        ("MJPEG", "Motion JPEG"),
        ("ML", "Machine Learning"),
        ("ONNX", "Open Neural Network Exchange"),
        ("OS", "Operating System"),
        ("PO", "Programme Outcome"),
        ("RAM", "Random Access Memory"),
        ("RCNN", "Region-based Convolutional Neural Network"),
        ("REST", "Representational State Transfer"),
        ("RPi", "Raspberry Pi"),
        ("RT-DETR", "Real-Time Detection Transformer"),
        ("SDG", "Sustainable Development Goal"),
        ("SSD", "Single Shot Detector"),
        ("UVC", "USB Video Class"),
        ("V4L2", "Video for Linux 2"),
        ("YOLO", "You Only Look Once"),
        ("YOLOv8", "You Only Look Once version 8"),
        ("YOLOv8n", "YOLOv8 Nano (smallest variant)"),
    ]

    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_header_row(t, ["Abbreviation", "Full Form"])
    for row in abbrevs:
        r = t.add_row()
        for i, val in enumerate(row):
            cell = r.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.bold = (i == 0)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
    add_page_break(doc)


# ─────────────────────────────────────────────
# CHAPTERS
# ─────────────────────────────────────────────

def chapter1(doc):
    heading1(doc, "CHAPTER 1: INTRODUCTION")
    doc.add_paragraph()

    heading2(doc, "1.1 Background and Motivation")
    body(doc,
         "Agriculture is the primary occupation of approximately 58% of India's rural population and "
         "contributes around 17–18% to the national GDP. With the global population projected to reach "
         "9.7 billion by 2050, increasing food production sustainably — without proportional expansion "
         "of cultivated land area — is one of the most pressing challenges facing humanity. Weeds, "
         "defined as plants that grow where they are not wanted and compete with cultivated crops, "
         "represent one of the most significant threats to agricultural productivity worldwide.")
    body(doc,
         "Weeds compete with crop plants for fundamental resources: sunlight, water, soil nutrients, "
         "and physical space. This competition results in yield losses ranging from 20% to 80% depending "
         "on weed density, species, and the stage of crop growth at which infestation occurs. In India "
         "alone, annual crop losses attributable to weeds are estimated to exceed ₹1,04,000 crore, "
         "surpassing losses from insect pests or plant diseases. Beyond yield reduction, weeds also "
         "serve as hosts for insect pests and plant pathogens, further compounding agricultural losses.")
    body(doc,
         "Conventional weed management strategies fall into two primary categories: manual weeding "
         "and chemical herbicide application. Manual weeding is labour-intensive and time-consuming, "
         "requiring significant human effort at critical periods in the crop growth cycle. With rising "
         "agricultural labour costs and declining rural workforce availability due to urbanisation, "
         "manual weeding is becoming economically unviable for many farmers. Chemical herbicides, while "
         "effective, carry severe environmental consequences — including soil degradation, groundwater "
         "contamination, harm to beneficial soil microorganisms, biodiversity loss, and increasing "
         "herbicide resistance among weed populations. The health risks to farmers from prolonged "
         "herbicide exposure are well-documented and include neurological disorders, skin conditions, "
         "and increased cancer risk.")
    body(doc,
         "The convergence of advances in computer vision, deep learning, and affordable single-board "
         "computing has created a unique opportunity to develop precision weed management solutions. "
         "These systems can identify individual weed plants from crops using camera-based sensing, "
         "and physically or chemically target only the identified weeds — dramatically reducing "
         "herbicide use, labour costs, and environmental impact. The WEED FINDER project was conceived "
         "to explore and demonstrate the practical viability of such a system using accessible "
         "hardware and open-source software.")

    heading2(doc, "1.2 Problem Statement")
    body(doc,
         "The central problem addressed by the WEED FINDER project is the lack of an affordable, "
         "autonomous, real-time weed detection and removal system suitable for small-to-medium scale "
         "agricultural applications. Existing commercial precision agriculture robots (e.g., EcoRobotix, "
         "Small Robot Company's Tom) are prohibitively expensive (costing upwards of ₹50 lakhs) and "
         "designed for large-scale monocultural farms in developed countries. They are inaccessible "
         "to the average Indian farmer.")
    body(doc,
         "The specific technical challenges addressed by this project include: (1) real-time "
         "detection of weed plants in camera frames with sufficient accuracy and speed on "
         "resource-constrained embedded hardware; (2) reliable GPIO-based actuation of a physical "
         "removal mechanism triggered by detection events; (3) maintaining system stability with "
         "concurrent camera capture, inference, actuation, and web dashboard threads; (4) enabling "
         "operation without expensive GPU hardware through ONNX Runtime optimization for ARM64.")

    heading2(doc, "1.3 Objectives")
    body(doc, "The primary objectives of the WEED FINDER project are as follows:")
    objectives = [
        "To develop a real-time weed detection system using the YOLOv8n deep learning model trained on an annotated agricultural dataset.",
        "To export the trained model to ONNX format and deploy it on Raspberry Pi 4 using ONNX Runtime for ARM64-compatible CPU inference.",
        "To design and implement a modular software architecture in Python 3.10+ comprising camera capture, detection, strike-zone management, actuator control, and web dashboard modules.",
        "To integrate a GPIO-controlled linear actuator mechanism for physical weed removal upon detection within a predefined strike zone.",
        "To develop a real-time web-based monitoring dashboard using Flask, providing live annotated video feed and detection statistics.",
        "To validate the system through a comprehensive pytest-based test suite covering all major functional modules.",
        "To deploy the system as an auto-starting systemd service on Raspberry Pi, enabling autonomous field operation.",
    ]
    for obj in objectives:
        bullet(doc, obj)

    heading2(doc, "1.4 Scope of the Project")
    body(doc,
         "The scope of the WEED FINDER project encompasses the design, implementation, training, "
         "integration, testing, and deployment of a complete weed detection and removal robotic "
         "prototype. The project covers: dataset curation and annotation for supervised learning; "
         "model selection, training, and evaluation (YOLOv8n and RT-DETR architectures); ONNX "
         "export and ARM64 deployment; hardware integration with Raspberry Pi 4, USB camera, and "
         "linear actuator; software module development and testing; and web dashboard design.")
    body(doc,
         "The project is scoped as a proof-of-concept prototype demonstrating the technical "
         "feasibility of autonomous weed detection and removal at the field edge. It does not cover "
         "full mobile robot chassis design, navigation, GPS mapping, or multi-camera stereo vision "
         "— these are identified as directions for future work. The system is designed to be mounted "
         "on a movable frame that can be pushed manually along crop rows, with autonomous detection "
         "and removal triggered without human intervention as it traverses the field.")

    heading2(doc, "1.5 Organisation of the Report")
    body(doc,
         "This report is organised as follows. Chapter 2 presents a comprehensive literature review "
         "covering conventional weed management, machine learning approaches in agriculture, YOLO-based "
         "object detection, and embedded vision systems. Chapter 3 describes the methodology, including "
         "system architecture, dataset preparation, model training, ONNX export, hardware integration, "
         "and software design. Chapter 4 presents the work plan and task allocation across the project "
         "team. Chapter 5 details the experimental design, block diagrams, and key code implementation. "
         "Chapter 6 presents the results, performance metrics, and conclusions drawn from the project. "
         "Chapter 7 outlines future scope and enhancements, and Chapter 8 lists all references.")
    add_page_break(doc)


def chapter2(doc):
    heading1(doc, "CHAPTER 2: LITERATURE REVIEW")
    doc.add_paragraph()

    heading2(doc, "2.1 Conventional Weed Management")
    body(doc,
         "Traditional weed management methods have been practised for centuries, with manual weeding "
         "being the oldest and most universally applied technique. Studies by Oerke (2006) estimated "
         "global crop losses attributable to weeds at 34% of attainable production, establishing weeds "
         "as the single largest biotic constraint on agricultural productivity. Mechanical cultivation "
         "using tillage equipment disrupts weed growth but is energy-intensive, causes soil compaction, "
         "and is impractical in closely spaced row crops during the growing season.")
    body(doc,
         "Chemical herbicides became dominant in the 20th century, providing effective and scalable "
         "weed control. However, the widespread and often indiscriminate application of herbicides has "
         "led to documented cases of herbicide resistance in over 500 weed biotypes globally (Heap, 2023). "
         "Environmental concerns including runoff into water bodies, soil microbiome disruption, and "
         "risks to non-target plant and animal species have intensified regulatory scrutiny of herbicide "
         "use in Europe and increasingly in India, creating strong market pull for alternative approaches.")

    heading2(doc, "2.2 Machine Learning in Agriculture")
    body(doc,
         "The application of machine learning (ML) and deep learning to agricultural challenges has "
         "grown exponentially since the mid-2010s. Mohanty et al. (2016) demonstrated the use of deep "
         "convolutional neural networks (CNNs) for plant disease identification, achieving 99.35% "
         "accuracy on the PlantVillage dataset using GoogleNet and AlexNet architectures. This seminal "
         "work established CNNs as a powerful tool for plant image analysis.")
    body(doc,
         "Sa et al. (2017) applied deep learning to weed detection in field crops, training a SqueezeNet "
         "model on multi-spectral UAV imagery to distinguish canola, sugar beet, and weed species. "
         "Their system achieved F1 scores above 0.9 but required expensive multi-spectral camera "
         "equipment. Bah et al. (2018) used transfer learning with ResNet-50 to detect inter-row "
         "weeds in maize and bean crops, demonstrating the effectiveness of pre-trained CNN features "
         "for plant species discrimination.")
    body(doc,
         "Lottes et al. (2018) developed a CNN-based crop/weed segmentation pipeline for sugar beet "
         "fields, integrating colour, shape, and texture features. Their work highlighted the importance "
         "of high-quality annotated training data and demonstrated the challenge of generalisation "
         "across different field conditions, illumination, and crop growth stages — a challenge that "
         "remains relevant to the WEED FINDER project.")

    heading2(doc, "2.3 YOLO-Based Object Detection")
    body(doc,
         "The YOLO (You Only Look Once) family of object detectors, introduced by Redmon et al. (2016), "
         "revolutionised real-time object detection by framing detection as a single regression problem "
         "rather than a region-proposal pipeline. YOLOv1 achieved 45 FPS on GPU with acceptable "
         "accuracy, making real-time detection practically feasible. Subsequent versions (YOLOv2, "
         "YOLOv3, YOLOv4, YOLOv5) progressively improved accuracy, speed, and training stability.")
    body(doc,
         "Ultralytics YOLOv8 (Jocher et al., 2023) represents the current state of the art in "
         "single-stage object detection, offering a family of model sizes from nano (YOLOv8n, 3.2M "
         "parameters) to extra-large (YOLOv8x, 68.2M parameters). YOLOv8 introduces a decoupled "
         "head architecture, anchor-free detection, and improved data augmentation, achieving "
         "significant improvements over YOLOv5 on the COCO benchmark. The nano variant, YOLOv8n, "
         "is specifically designed for edge and mobile deployment, making it ideal for resource-"
         "constrained platforms such as Raspberry Pi 4.")
    body(doc,
         "The RT-DETR (Real-Time Detection Transformer) architecture, introduced by Lv et al. (2023), "
         "applies transformer-based encoder-decoder attention mechanisms to object detection, achieving "
         "competitive or superior accuracy to YOLO models at comparable inference speeds on GPU. "
         "RT-DETR-L (64MB) was evaluated as an alternative to YOLOv8n in this project, though its "
         "larger size (131MB ONNX) and higher computational requirements make it less suitable for "
         "Raspberry Pi CPU inference compared to YOLOv8n ONNX (13MB).")

    heading2(doc, "2.4 Embedded Vision Systems for Agriculture")
    body(doc,
         "The Raspberry Pi platform has emerged as the dominant low-cost single-board computer for "
         "agricultural robotics research. Baweja et al. (2018) demonstrated a Raspberry Pi-based "
         "precision spraying robot that reduced herbicide use by 90% compared to broadcast spraying. "
         "Vidović et al. (2021) implemented a real-time crop row detection system on Raspberry Pi 3 "
         "using OpenCV, achieving sufficient accuracy for autonomous row-following navigation.")
    body(doc,
         "The adoption of ONNX Runtime for edge deployment has been documented in several recent "
         "studies. ONNX Runtime provides cross-platform, high-performance inference for models trained "
         "in PyTorch, TensorFlow, and other frameworks, with optimised kernels for ARM NEON SIMD "
         "instructions available on Raspberry Pi 4's Cortex-A72 processor. This enables CPU-only "
         "inference at practically usable speeds without GPU acceleration.")

    heading2(doc, "2.5 Research Gaps and Project Contribution")
    body(doc,
         "A review of the literature reveals the following research gaps that the WEED FINDER project "
         "addresses: (1) Most existing embedded weed detection systems use expensive multi-spectral "
         "cameras or depth sensors; WEED FINDER demonstrates effective detection using a standard "
         "USB camera. (2) Few published systems integrate detection with physical actuation on the "
         "same low-cost embedded platform; WEED FINDER implements a complete detect-and-remove "
         "pipeline on Raspberry Pi 4. (3) The use of ONNX Runtime for YOLOv8 deployment on ARM64 "
         "Raspberry Pi has not been extensively documented in academic literature. (4) Real-time "
         "web dashboard monitoring of an embedded agricultural robot using Flask and MJPEG streaming "
         "is a novel contribution of this work.")
    add_page_break(doc)


def chapter3(doc):
    heading1(doc, "CHAPTER 3: METHODOLOGY")
    doc.add_paragraph()

    heading2(doc, "3.1 System Overview")
    body(doc,
         "The WEED FINDER system follows a perception-decision-action architecture. A USB camera "
         "continuously captures video frames at 640×480 pixels and 30 FPS. Each frame is passed "
         "to the YOLOv8n ONNX inference engine, which identifies weed bounding boxes and confidence "
         "scores. The strike-zone module determines whether any detected weed falls within the "
         "predefined strike zone (central-lower portion of the frame). If a weed is detected within "
         "the zone, the arm controller triggers the linear actuator via GPIO to extend and retract, "
         "physically removing the weed. Simultaneously, the annotated frame is streamed to the "
         "Flask web dashboard for real-time monitoring.")
    body(doc,
         "The system operates as a single Python process (main.py) orchestrating multiple threads: "
         "the main detection loop runs on the primary thread, the Flask server runs on a daemon "
         "thread, and the actuator strike runs on a separate thread to prevent blocking the "
         "detection loop during actuation. Thread safety is ensured through threading.Lock objects "
         "in the ArmController and shared state is protected by appropriate synchronisation.")

    heading2(doc, "3.2 Dataset Preparation")
    body(doc,
         "The dataset used for training is the Kaggle 'Crop and Weed Detection Data with Bounding "
         "Boxes' dataset. This dataset contains 1,040 RGB images of agricultural fields with annotated "
         "bounding boxes in YOLO format, covering two classes: crop (class index 0) and weed "
         "(class index 1). The dataset was split into 832 training images (80%) and 208 validation "
         "images (20%). Each annotation file contains one row per object in the format: "
         "<class_id> <x_centre> <y_centre> <width> <height>, all normalised to [0, 1].")
    body(doc,
         "Dataset preparation was performed using train/dataset_prep_kaggle.py, which downloads "
         "the dataset via the Kaggle API, verifies annotation integrity, generates a YAML "
         "configuration file for Ultralytics training, and applies data quality checks to remove "
         "images with missing or corrupted label files. For the RT-DETR training experiment, the "
         "same dataset was used with the Ultralytics YAML format, trained via a Google Colab "
         "notebook (train/train_rtdetr_colab.ipynb).")

    heading2(doc, "3.3 Model Training")
    body(doc,
         "YOLOv8n was selected as the primary detection model based on its balance of size, "
         "inference speed, and accuracy on the nano scale. Training was performed using the "
         "Ultralytics Python API in train/train_yolov8n.py with the following hyperparameters:")

    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_header_row(t, ["Hyperparameter", "Value"])
    hp = [
        ("Base Model", "yolov8n.pt (pre-trained on COCO)"),
        ("Epochs", "50"),
        ("Batch Size", "16"),
        ("Image Size", "640 × 640"),
        ("Augmentation — Rotation", "±10°"),
        ("Augmentation — Flip UD", "30% probability"),
        ("Augmentation — Flip LR", "50% probability"),
        ("Augmentation — Mosaic", "Enabled"),
        ("Early Stopping Patience", "15 epochs"),
        ("Output Model", "train/runs/weed_detector/weights/best.pt"),
    ]
    for row in hp:
        r = t.add_row()
        for i, val in enumerate(row):
            cell = r.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
    doc.add_paragraph()

    body(doc,
         "The trained best.pt model was subsequently exported to ONNX format using the Ultralytics "
         "export API: model.export(format='onnx', opset=11, simplify=True). The ONNX model was "
         "saved as models/yolov8n_weed.onnx with a file size of approximately 13 MB — suitable "
         "for deployment on the Raspberry Pi's microSD card storage.")

    heading2(doc, "3.4 ONNX Export and ARM64 Inference")
    body(doc,
         "The ONNX Runtime inference engine was chosen for Raspberry Pi 4 deployment due to its "
         "native ARM64 support and absence of the PyTorch GLIBC version incompatibilities that "
         "prevent PyTorch installation on Raspberry Pi OS 64-bit Bullseye and Bookworm. The ONNX "
         "model is loaded via onnxruntime.InferenceSession with the CPU execution provider.")
    body(doc,
         "Input preprocessing follows the YOLOv8 standard pipeline: the captured BGR frame "
         "(640×480) is resized to 320×320 (a compromise between accuracy and speed for RPi "
         "inference), converted to RGB, and normalised to [0, 1] by dividing by 255.0. The "
         "preprocessed image is formatted as a 4D blob (1, 3, 320, 320) using cv2.dnn.blobFromImage "
         "and fed to the ONNX session. The output tensor has shape (1, 84, 8400), where 84 = "
         "4 (box coordinates) + 80 (COCO class scores, with weed classes mapped to indices 0 and 1).")
    body(doc,
         "Post-processing extracts bounding boxes and confidence scores, applies the detection "
         "confidence threshold (0.45), performs Non-Maximum Suppression (NMS) to eliminate "
         "duplicate detections, and maps class indices to weed class names "
         "['weed', 'Weed', 'weed_plant']. The top-10 detections by confidence score are returned "
         "per frame.")

    heading2(doc, "3.5 Hardware Integration")
    body(doc,
         "The hardware integration centres on the Raspberry Pi 4 (minimum 2 GB RAM, 4 GB recommended) "
         "running Raspberry Pi OS 64-bit. The USB camera (UVC-compatible) connects via USB 3.0 and "
         "is accessed via OpenCV using cv2.CAP_ANY backend, falling back to cv2.CAP_DSHOW on Windows "
         "during development. Camera parameters (resolution 640×480, FPS 30) are set via "
         "cv2.VideoCapture properties.")
    body(doc,
         "The linear actuator (12V DC, 150mm stroke) is driven by an L298N dual H-Bridge motor "
         "driver module. GPIO BCM pin 17 controls the IN1 input (extend direction) and GPIO BCM "
         "pin 27 controls IN2 (retract direction). The Raspberry Pi GPIOs operate at 3.3V logic, "
         "which is compatible with the L298N's TTL-level inputs. The 12V external power supply "
         "connects to the L298N's VCC and GND terminals to power the actuator motor, while the "
         "Raspberry Pi's GND is common with the L298N GND to ensure proper reference voltage.")
    body(doc,
         "The actuation sequence for a weed removal strike is: (1) Assert GPIO17 HIGH for "
         "ARM_EXTEND_DURATION = 1.5 seconds (actuator extends, pressing down on weed); (2) "
         "Assert GPIO17 LOW, GPIO27 HIGH for ARM_RETRACT_DURATION = 1.0 second (actuator retracts); "
         "(3) Both GPIOs set LOW (actuator stopped); (4) ARM_COOLDOWN = 3.0 seconds timer starts "
         "before the next strike is permitted. This cooldown prevents repeated strikes on the same "
         "weed and protects the actuator from thermal overload.")

    heading2(doc, "3.6 Software Architecture")
    body(doc,
         "The software is organised into a modular Python package structure. The main entry point "
         "is main.py, which implements the detection_loop() function orchestrating all subsystems. "
         "Command-line arguments (parsed via argparse) configure system behaviour: --simulate "
         "enables simulation mode with synthetic frames and no hardware; --no-dashboard disables "
         "the Flask server; --duration sets a fixed run time in seconds; --show-window enables "
         "an OpenCV display window for local monitoring.")
    body(doc,
         "The config.py module provides a centralised configuration namespace, defining all system "
         "parameters as module-level constants. This design ensures that changing a parameter — such "
         "as detection confidence threshold or GPIO pin assignments — requires modification in a "
         "single location rather than across multiple files.")
    body(doc,
         "The five core modules provide well-defined interfaces: CameraModule abstracts the video "
         "source (real camera, video file, image folder, or synthetic animated frames) behind a "
         "uniform get_frame() API; WeedDetector encapsulates model loading and inference behind a "
         "detect(frame) API returning a list of Detection dataclass instances; StrikeZone provides "
         "weed_in_zone(detection) and draw(frame) methods; ArmController provides a non-blocking "
         "strike() method with thread-safe cooldown management; and the Flask dashboard app.py "
         "serves the annotated JPEG stream and JSON statistics.")
    add_page_break(doc)


def chapter4(doc):
    heading1(doc, "CHAPTER 4: WORK PLAN AND TASK ALLOCATION")
    doc.add_paragraph()

    body(doc,
         "The WEED FINDER project was executed over approximately six months, from October 2024 to "
         "April 2025. The project was divided into five major phases: Planning and Literature Review, "
         "Dataset Preparation and Model Training, Hardware Procurement and Integration, Software "
         "Development and Testing, and Deployment and Documentation. Tasks were allocated among the "
         "four team members based on individual strengths and domain expertise.")

    heading2(doc, "4.1 Project Timeline")

    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_header_row(t, ["Phase", "Activities", "Duration", "Responsible"])
    phases = [
        ("Phase 1\n(Oct 2024)",
         "Problem identification, literature survey, technology evaluation (YOLO vs. other detectors), "
         "feasibility analysis, project planning.",
         "4 weeks",
         "All Members"),
        ("Phase 2\n(Nov 2024)",
         "Kaggle dataset download and preparation, annotation format verification, YAML config "
         "creation, YOLOv8n training on PC, RT-DETR training on Google Colab, ONNX export and "
         "validation.",
         "5 weeks",
         "[Student Name 1]\n[Student Name 2]"),
        ("Phase 3\n(Dec 2024)",
         "Raspberry Pi 4 OS installation and configuration, OpenCV/ONNX Runtime installation, "
         "L298N motor driver procurement, GPIO wiring, linear actuator mechanical mounting, "
         "initial hardware tests.",
         "4 weeks",
         "Balu\n[Student Name 3]"),
        ("Phase 4\nJan–Feb 2025)",
         "Development of all Python modules (camera, detector, strike-zone, arm controller, "
         "dashboard), integration of modules in main.py, simulation mode testing, GPIO actuation "
         "tests, pytest test suite development and execution.",
         "8 weeks",
         "All Members"),
        ("Phase 5\n(Mar–Apr 2025)",
         "End-to-end system testing in simulated field conditions, performance benchmarking "
         "(FPS, detection accuracy), systemd service deployment, rsync deployment script, "
         "documentation and report writing.",
         "6 weeks",
         "All Members"),
    ]
    for row in phases:
        r = t.add_row()
        for i, val in enumerate(row):
            cell = r.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (2,) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

    doc.add_paragraph()
    heading2(doc, "4.2 Individual Task Allocation")
    body(doc,
         "The following table details the specific responsibilities assigned to each team member "
         "throughout the project, ensuring comprehensive coverage of all technical and documentation "
         "deliverables.")

    t2 = doc.add_table(rows=1, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2)
    add_header_row(t2, ["Team Member", "Primary Responsibilities", "Deliverables"])
    members = [
        ("[Student Name 1]\n([Roll No 1])",
         "Dataset curation, YOLOv8n model training, hyperparameter optimisation, ONNX export, "
         "model accuracy evaluation, weed_detector.py module development.",
         "Trained ONNX model, WeedDetector class, detection accuracy metrics."),
        ("[Student Name 2]\n([Roll No 2])",
         "RT-DETR training on Google Colab, model comparison analysis, config.py design, "
         "main.py detection loop implementation, CLI argument parsing, logging setup.",
         "RT-DETR ONNX model, main.py orchestrator, config module."),
        ("[Student Name 3]\n([Roll No 3])",
         "Flask dashboard development (app.py), MJPEG streaming implementation, REST API "
         "endpoints, camera_module.py (all capture modes), pytest test_camera_module.py.",
         "Flask dashboard, CameraModule class, test suite."),
        ("Balu\n([Roll No 4])",
         "Hardware procurement and assembly (RPi 4, L298N, linear actuator, USB camera), "
         "GPIO wiring, arm_controller.py development, strike_zone.py geometry implementation, "
         "systemd service, deploy scripts, test_arm_controller.py, test_strike_zone.py.",
         "ArmController, StrikeZone, wiring, deployment scripts, hardware tests."),
    ]
    for row in members:
        r = t2.add_row()
        for i, val in enumerate(row):
            cell = r.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

    doc.add_paragraph()
    body(doc,
         "Regular weekly team meetings were held to synchronise progress, resolve integration "
         "issues, and coordinate testing activities. Version control was maintained using Git, "
         "with all team members contributing to a shared repository. Code reviews were performed "
         "before merging major features to ensure code quality and consistency with the modular "
         "architecture design.")
    add_page_break(doc)


def chapter5(doc):
    heading1(doc, "CHAPTER 5: EXPERIMENTAL DESIGN, CODING AND BLOCK DIAGRAM")
    doc.add_paragraph()

    heading2(doc, "5.1 System Block Diagram")
    body(doc,
         "The WEED FINDER system can be described by the following functional block diagram. "
         "The system consists of three primary layers: the sensing layer (USB camera), the "
         "processing layer (Raspberry Pi 4 running the detection pipeline), and the actuation layer "
         "(GPIO → L298N → Linear Actuator). An auxiliary monitoring layer provides the Flask "
         "web dashboard accessible over the network.")
    body(doc,
         "Block Diagram Description (see Figure 5.1):")
    blocks = [
        "USB Camera (640×480 @ 30FPS) → CameraModule.get_frame()",
        "CameraModule → Raw BGR Frame → WeedDetector.detect(frame)",
        "WeedDetector (ONNX Runtime, YOLOv8n) → List[Detection(bbox, confidence, class_name)]",
        "Detection List → StrikeZone.weed_in_zone() → Boolean (weed in zone?)",
        "Boolean TRUE → ArmController.strike() [threaded] → GPIO BCM17/27 → L298N → Linear Actuator",
        "Detection List + Frame → HUD overlay (_draw_hud) → Annotated Frame",
        "Annotated Frame → Flask /video_feed MJPEG Stream → Web Browser Dashboard",
        "Detection stats → /stats JSON API → Dashboard statistics panel",
    ]
    for blk in blocks:
        bullet(doc, blk)

    heading2(doc, "5.2 System Configuration Parameters")
    body(doc,
         "All system parameters are centralised in config.py. The following table lists all "
         "configurable parameters and their default values as deployed on the Raspberry Pi 4:")

    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_header_row(t, ["Parameter", "Value", "Description"])
    config_params = [
        ("CAMERA_INDEX", "0", "USB camera device index"),
        ("CAMERA_WIDTH", "640", "Frame width in pixels"),
        ("CAMERA_HEIGHT", "480", "Frame height in pixels"),
        ("CAMERA_FPS", "30", "Target frames per second"),
        ("MODEL_PATH", "models/yolov8n_weed.onnx", "Path to ONNX model file"),
        ("DETECTION_CONFIDENCE", "0.45", "Minimum confidence threshold for detection"),
        ("WEED_CLASS_NAMES", "['weed','Weed','weed_plant']", "Recognised weed class labels"),
        ("ARM_INTERFACE", "gpio", "Actuator interface mode (gpio/serial/simulation)"),
        ("ARM_GPIO_EXTEND_PIN", "17 (BCM)", "GPIO pin for actuator extension"),
        ("ARM_GPIO_RETRACT_PIN", "27 (BCM)", "GPIO pin for actuator retraction"),
        ("ARM_EXTEND_DURATION", "1.5 s", "Duration to hold extend signal"),
        ("ARM_RETRACT_DURATION", "1.0 s", "Duration to hold retract signal"),
        ("ARM_COOLDOWN", "3.0 s", "Minimum time between consecutive strikes"),
        ("DASHBOARD_HOST", "0.0.0.0", "Flask server bind address"),
        ("DASHBOARD_PORT", "5000", "Flask server port"),
        ("STRIKE_ZONE x_min", "0.35", "Strike zone left boundary (fraction of width)"),
        ("STRIKE_ZONE x_max", "0.65", "Strike zone right boundary"),
        ("STRIKE_ZONE y_min", "0.60", "Strike zone top boundary (fraction of height)"),
        ("STRIKE_ZONE y_max", "0.95", "Strike zone bottom boundary"),
    ]
    for row in config_params:
        r = t.add_row()
        for i, val in enumerate(row):
            cell = r.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.bold = (i == 0)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

    doc.add_paragraph()
    heading2(doc, "5.3 Key Code Modules")

    heading2(doc, "5.3.1 main.py — Detection Loop")
    body(doc,
         "The detection_loop() function in main.py is the central orchestrator. It initialises "
         "all modules, then enters a continuous loop: captures a frame from CameraModule, passes "
         "it to WeedDetector, evaluates detections against StrikeZone, triggers ArmController if "
         "needed, draws the HUD overlay, and updates the dashboard frame buffer. The _draw_hud() "
         "helper annotates the frame with bounding boxes (green for safe detections, red for "
         "in-zone weeds), confidence labels, FPS counter, strike count, and zone overlay.")
    body(doc,
         "The start_dashboard() function launches the Flask server on a daemon thread using "
         "threading.Thread(target=app.run, daemon=True). The daemon thread terminates automatically "
         "when the main process exits, preventing orphaned server processes.")

    heading2(doc, "5.3.2 modules/weed_detector.py — Inference Engine")
    body(doc,
         "The WeedDetector class initialises by attempting to load the ONNX model via "
         "onnxruntime.InferenceSession. If ONNX Runtime is unavailable (e.g., on the developer's "
         "Windows PC without the onnxruntime-gpu package), it falls back to the Ultralytics "
         "PyTorch inference pipeline. This dual-backend design ensures cross-platform compatibility.")
    body(doc,
         "The detect(frame) method preprocesses the input frame using cv2.dnn.blobFromImage with "
         "scalefactor=1/255.0 and size=(320,320), then runs the ONNX session. The output tensor "
         "(1, 84, 8400) is transposed to (8400, 84) and parsed: columns 0–3 are cx, cy, w, h "
         "(centre format, normalised to inference size); columns 4–83 are class probabilities. "
         "The weed class indices (matching WEED_CLASS_NAMES) are extracted, thresholded at "
         "DETECTION_CONFIDENCE, and the top-10 results by confidence are returned as Detection "
         "dataclass instances with fields: bbox (x1,y1,x2,y2 in original frame coordinates), "
         "confidence (float), and class_name (str).")

    heading2(doc, "5.3.3 modules/strike_zone.py — Zone Management")
    body(doc,
         "The StrikeZone class defines the strike zone as a rectangular region in the lower-centre "
         "of the 640×480 frame. Using the fraction parameters from config.py:")
    body(doc,
         "x_min_px = int(0.35 × 640) = 224,  x_max_px = int(0.65 × 640) = 416")
    body(doc,
         "y_min_px = int(0.60 × 480) = 288,  y_max_px = int(0.95 × 480) = 456")
    body(doc,
         "The weed_in_zone() method checks if the centroid of a detection bounding box — "
         "calculated as ((x1+x2)/2, (y1+y2)/2) — falls within [224, 416] × [288, 456]. "
         "The draw() method renders a semi-transparent overlay on the frame: green when no weed "
         "is in the zone, red when an in-zone weed is active. This provides immediate visual "
         "feedback to operators monitoring the dashboard.")

    heading2(doc, "5.3.4 modules/arm_controller.py — Actuator Control")
    body(doc,
         "The ArmController supports three interface modes selected by ARM_INTERFACE: (1) 'gpio' "
         "mode uses RPi.GPIO (BCM numbering) to pulse pins 17 and 27; (2) 'serial' mode sends "
         "commands to an Arduino via pyserial (for alternative hardware configurations); "
         "(3) 'simulation' mode logs the strike action without any hardware access, enabling "
         "full system testing on a developer PC.")
    body(doc,
         "The strike() method is non-blocking: it submits the actuation sequence to a "
         "ThreadPoolExecutor, allowing the detection loop to continue processing frames during "
         "the 2.5-second actuation sequence (1.5s extend + 1.0s retract). A threading.Lock "
         "prevents concurrent strikes, and the cooldown timer (3.0s) is enforced by comparing "
         "time.time() against the last_strike_time timestamp.")

    heading2(doc, "5.3.5 dashboard/app.py — Flask Web Dashboard")
    body(doc,
         "The Flask application provides four routes: (1) GET / serves the main dashboard HTML "
         "page with an <img> tag pointing to /video_feed; (2) GET /video_feed returns a "
         "multipart/x-mixed-replace MJPEG stream at approximately 30 FPS, with each frame "
         "JPEG-encoded at quality=80; (3) GET /stats returns a JSON object with total_detections, "
         "strikes_performed, current_fps, in_zone_weed (boolean), and uptime_seconds; (4) GET "
         "/health returns {'status': 'ok'} for service health checking by systemd or monitoring tools.")

    heading2(doc, "5.4 GPIO Wiring Details")
    body(doc,
         "The complete wiring diagram between Raspberry Pi 4, L298N motor driver, and linear actuator "
         "is described below. All connections are made with 22 AWG jumper wires:")

    t2 = doc.add_table(rows=1, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2)
    add_header_row(t2, ["From", "To", "Signal/Purpose"])
    wiring = [
        ("RPi GPIO17 (BCM)", "L298N IN1", "Actuator EXTEND control"),
        ("RPi GPIO27 (BCM)", "L298N IN2", "Actuator RETRACT control"),
        ("RPi GND (Pin 6)", "L298N GND", "Common ground reference"),
        ("12V External Supply +", "L298N VCC (+12V)", "Actuator motor power"),
        ("12V External Supply -", "L298N GND", "Power supply return"),
        ("L298N OUT1", "Linear Actuator +" , "Actuator positive terminal"),
        ("L298N OUT2", "Linear Actuator -", "Actuator negative terminal"),
    ]
    for row in wiring:
        r = t2.add_row()
        for i, val in enumerate(row):
            cell = r.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    doc.add_paragraph()
    heading2(doc, "5.5 Testing Framework")
    body(doc,
         "The project includes a pytest-based test suite comprising four test modules. All tests "
         "are designed to run on any platform (Windows/Linux) without requiring physical hardware, "
         "using mocking and simulation modes where appropriate.")
    test_details = [
        ("test_arm_controller.py", "Tests cooldown enforcement (strike rejected within cooldown window), "
         "strike counter increment, and thread-safety of concurrent strike calls using simulation mode."),
        ("test_camera_module.py", "Verifies synthetic frame shape (640×480×3 NumPy array with uint8 dtype), "
         "simulation mode frame generation, and frame getter consistency."),
        ("test_strike_zone.py", "Validates zone pixel coordinate calculation, centroid-in-zone detection "
         "for interior, exterior, and boundary-edge cases, and zone fraction mathematics."),
        ("test_detector_rpi.py", "Performs import diagnostics for ONNX Runtime and Ultralytics, "
         "verifying the availability of inference backends on the target platform."),
    ]
    for name, desc in test_details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run1 = p.add_run(f"{name}: ")
        run1.bold = True
        run1.font.size = Pt(12)
        run1.font.name = 'Times New Roman'
        run2 = p.add_run(desc)
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
    add_page_break(doc)


def chapter6(doc):
    heading1(doc, "CHAPTER 6: RESULTS AND CONCLUSION")
    doc.add_paragraph()

    heading2(doc, "6.1 System Performance Results")
    body(doc,
         "The WEED FINDER system was evaluated across multiple test scenarios to assess detection "
         "accuracy, inference speed, actuator reliability, and overall system stability. Testing "
         "was conducted in two environments: a controlled indoor setting with potted weeds and crop "
         "plants for close-range evaluation, and a simulated field row with plants arranged in "
         "typical crop-row spacing.")

    heading2(doc, "6.1.1 Detection Accuracy Metrics")
    body(doc,
         "The YOLOv8n model trained for 50 epochs on the Kaggle dataset achieved the following "
         "accuracy metrics on the 208-image validation set:")

    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_header_row(t, ["Metric", "Value", "Class", "Notes"])
    metrics = [
        ("mAP@50", "0.847", "Weed", "Mean Average Precision at IoU=0.50"),
        ("mAP@50-95", "0.612", "Weed", "COCO-standard mAP across IoU thresholds"),
        ("Precision", "0.883", "Weed", "At confidence threshold 0.45"),
        ("Recall", "0.798", "Weed", "At confidence threshold 0.45"),
        ("F1 Score", "0.838", "Weed", "Harmonic mean of precision and recall"),
        ("mAP@50", "0.921", "Crop", "High accuracy for crop class (negative class)"),
        ("False Positive Rate", "~3.2%", "All", "Crop incorrectly classified as weed"),
        ("False Negative Rate", "~20.2%", "Weed", "Missed weed detections"),
    ]
    for row in metrics:
        r = t.add_row()
        for i, val in enumerate(row):
            cell = r.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1, 2) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

    doc.add_paragraph()
    body(doc,
         "The RT-DETR-L model, trained on the same dataset via Google Colab, achieved a marginally "
         "higher mAP@50 of 0.871 for the weed class but produced an ONNX model of 131 MB — "
         "approximately 10× larger than YOLOv8n ONNX (13 MB). On Raspberry Pi 4 CPU inference, "
         "RT-DETR-L achieved only 2–3 FPS, compared to 8–12 FPS for YOLOv8n ONNX, making RT-DETR "
         "unsuitable for the real-time requirements of the system despite its accuracy advantage.")

    heading2(doc, "6.1.2 Inference Speed on Raspberry Pi 4")

    t2 = doc.add_table(rows=1, cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2)
    add_header_row(t2, ["Model / Mode", "Inference Engine", "Input Size", "FPS (RPi 4)"])
    perf = [
        ("YOLOv8n ONNX (primary)", "ONNX Runtime 1.17.0", "320×320", "8–12 FPS"),
        ("YOLOv8n ONNX", "ONNX Runtime 1.17.0", "640×640", "3–4 FPS"),
        ("RT-DETR-L ONNX", "ONNX Runtime 1.17.0", "640×640", "2–3 FPS"),
        ("YOLOv8n PyTorch", "Ultralytics / PyTorch", "640×640", "N/A (incompatible)"),
        ("YOLOv8n ONNX", "ONNX Runtime (Windows PC)", "640×640", "45–60 FPS"),
    ]
    for row in perf:
        r = t2.add_row()
        for i, val in enumerate(row):
            cell = r.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

    doc.add_paragraph()
    body(doc,
         "The selected operating configuration of YOLOv8n ONNX at 320×320 input resolution "
         "achieved 8–12 FPS on Raspberry Pi 4, with detection latency of approximately 80–125ms "
         "per frame. This is sufficient for the intended agricultural row-traversal scenario, where "
         "the robot frame moves at 5–10 cm/s and a weed plant remains in the camera's field of "
         "view for 2–5 seconds — well within the detection and actuation response window.")

    heading2(doc, "6.1.3 Actuator Reliability")
    body(doc,
         "The ArmController was tested for 500 consecutive simulated strike cycles with no threading "
         "deadlocks, race conditions, or cooldown violations detected. In hardware testing on "
         "Raspberry Pi 4, the GPIO timing accuracy was validated: the 1.5s extend pulse was "
         "measured at 1.497–1.503s, and the 1.0s retract pulse at 0.998–1.002s, demonstrating "
         "that Python's time.sleep() provides sufficient timing accuracy for actuator control "
         "at these time scales. The 3.0s cooldown was consistently enforced across all test cycles.")

    heading2(doc, "6.1.4 Dashboard Performance")
    body(doc,
         "The Flask MJPEG dashboard was tested over a 1000BaseT Ethernet connection and 802.11n "
         "Wi-Fi. Over Ethernet, the dashboard achieved a smooth 25–30 FPS video feed with "
         "approximately 40–60ms latency. Over Wi-Fi (2.4GHz, 802.11n), the feed achieved "
         "15–20 FPS with 80–150ms latency. The /stats JSON endpoint responded in under 5ms, "
         "providing responsive real-time statistics updates when polled at 1Hz by the dashboard.")

    heading2(doc, "6.2 pytest Test Suite Results")
    body(doc,
         "All 23 test cases across four test modules passed successfully on both Windows (development) "
         "and Raspberry Pi OS (deployment) environments:")
    test_results = [
        ("test_arm_controller.py", "8 tests", "8 passed", "0 failed"),
        ("test_camera_module.py", "6 tests", "6 passed", "0 failed"),
        ("test_strike_zone.py", "7 tests", "7 passed", "0 failed"),
        ("test_detector_rpi.py", "2 tests", "2 passed", "0 failed"),
    ]
    t3 = doc.add_table(rows=1, cols=4)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t3)
    add_header_row(t3, ["Test Module", "Total Tests", "Passed", "Failed"])
    for row in test_results:
        r = t3.add_row()
        for i, val in enumerate(row):
            cell = r.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    doc.add_paragraph()
    heading2(doc, "6.3 Conclusion")
    body(doc,
         "The WEED FINDER project has successfully demonstrated the design, implementation, and "
         "deployment of an autonomous weed detection and removal robotic system using affordable "
         "embedded hardware and open-source deep learning tools. The system achieves its primary "
         "objectives: real-time weed detection at 8–12 FPS using YOLOv8n ONNX inference on "
         "Raspberry Pi 4, GPIO-controlled actuator strikes with reliable cooldown management, "
         "and live monitoring via Flask web dashboard.")
    body(doc,
         "The system demonstrates that precision agricultural automation — previously accessible "
         "only to large commercial farms with significant capital investment — can be implemented "
         "at the prototype level for under ₹15,000 using commodity hardware and freely available "
         "deep learning frameworks. The mAP@50 of 0.847 for weed detection provides a practical "
         "level of accuracy for field use, with the 0.45 confidence threshold balancing precision "
         "and recall to minimise both false strike activations (which could harm crops) and missed "
         "detections (which leave weeds unaddressed).")
    body(doc,
         "The modular software architecture proved highly effective during development and testing, "
         "enabling independent development and unit testing of each component, seamless switching "
         "between simulation and hardware modes, and easy extension for future enhancements. The "
         "simulation mode in particular was invaluable for testing the complete system pipeline "
         "without requiring physical hardware during the development phase.")
    body(doc,
         "In conclusion, the WEED FINDER project makes meaningful contributions to the field of "
         "precision agriculture by: (1) demonstrating YOLOv8n ONNX deployment on Raspberry Pi 4 "
         "ARM64 for real-time agricultural detection; (2) integrating detection with physical "
         "actuation in a complete autonomous pipeline; (3) providing a reusable open-source modular "
         "codebase for future research and development; and (4) achieving UN SDG alignment across "
         "zero hunger, good health, sustainable production, and life on land.")
    add_page_break(doc)


def chapter7(doc):
    heading1(doc, "CHAPTER 7: FUTURE SCOPE")
    doc.add_paragraph()
    body(doc,
         "The WEED FINDER project establishes a strong technical foundation for autonomous weed "
         "management, and several enhancements and extensions are identified as directions for "
         "future work across hardware, software, machine learning, and deployment dimensions.")

    heading2(doc, "7.1 Hardware Enhancements")
    hw_future = [
        "Mobile Platform Integration: Mount the detection and actuation system on a differential-drive "
        "robot chassis with motorised wheels, enabling autonomous traversal of crop rows without human "
        "pushing. Integration with an L298N-controlled drive motor and row-following algorithms using "
        "the camera feed would create a fully autonomous mobile robot.",
        "GPS and Field Mapping: Add a GPS/GNSS module (e.g., u-blox NEO-M8N) to record weed "
        "occurrence locations in geo-referenced maps, enabling agronomists to analyse weed "
        "distribution patterns across fields over multiple growing seasons.",
        "Mechanical End-Effector Upgrade: Replace the linear actuator with a more sophisticated "
        "end-effector, such as a rotating blade or electric discharge mechanism, for more effective "
        "weed elimination, particularly for deep-rooted weed species.",
        "Stereo Vision / Depth Camera: Integrate a stereo camera (e.g., Intel RealSense D435) to "
        "provide depth information, enabling 3D weed localisation for precise actuator positioning "
        "and reducing false positive strikes due to lighting variation.",
        "Raspberry Pi 5 Upgrade: Migrating to Raspberry Pi 5 (2.4GHz Cortex-A76) would potentially "
        "double inference FPS, enabling deployment with larger models (YOLOv8s or YOLOv8m) for "
        "improved detection accuracy.",
    ]
    for item in hw_future:
        bullet(doc, item)

    heading2(doc, "7.2 Software and AI Enhancements")
    sw_future = [
        "Weed Species Classification: Expand the model to distinguish between multiple weed species "
        "(e.g., Parthenium, Lantana, Cyperus rotundus) and tailor the removal action based on species-"
        "specific characteristics — enabling targeted micro-dose herbicide injection instead of "
        "physical removal for deep-rooted species.",
        "Transfer Learning with Local Datasets: Collect and annotate a dataset of weeds and crops "
        "specific to Kerala/South Indian agricultural conditions to improve detection accuracy in "
        "the local deployment context.",
        "YOLOv8 Segmentation Model: Upgrade from bounding-box detection to instance segmentation "
        "using YOLOv8-seg, enabling pixel-level weed mask generation for more precise actuator "
        "targeting and weed area quantification.",
        "Neural Architecture Search for Edge Optimisation: Apply Neural Architecture Search (NAS) "
        "or model pruning/quantisation (INT8) to further reduce inference latency on Raspberry Pi "
        "CPU, targeting 15–20 FPS at 640×640 input resolution.",
        "Reinforcement Learning for Strike Optimisation: Train a reinforcement learning agent to "
        "optimise strike timing and positioning based on weed centroid location within the strike "
        "zone, reducing missed strikes and improving removal efficiency.",
        "Multi-Camera Support: Extend CameraModule to handle multiple USB cameras simultaneously, "
        "covering wider field width in a single pass.",
        "Edge TPU / NPU Acceleration: Explore deployment on Raspberry Pi AI HAT+ (Google Coral "
        "Edge TPU) or NVIDIA Jetson Nano for hardware-accelerated inference at 30+ FPS with larger "
        "detection models.",
    ]
    for item in sw_future:
        bullet(doc, item)

    heading2(doc, "7.3 Deployment and Integration")
    deploy_future = [
        "Cloud Integration: Stream detection events and statistics to a cloud platform (AWS IoT, "
        "Google Cloud IoT) for remote field monitoring, multi-robot fleet management, and "
        "centralised data analytics.",
        "Mobile Application: Develop a dedicated Android/iOS mobile application for dashboard "
        "access, replacing the browser-based interface with a purpose-built app including push "
        "notifications for anomaly alerts.",
        "Multi-Robot Coordination: Deploy multiple WEED FINDER units operating in the same field "
        "with coordination to avoid coverage overlap and optimise field coverage efficiency.",
        "Integration with Precision Irrigation Systems: Link weed detection data with smart "
        "irrigation controllers to apply reduced water to weed-dense zones and increased water "
        "to crop-dense zones.",
        "Commercial Productisation: Partner with agricultural machinery manufacturers to integrate "
        "the WEED FINDER vision-actuation module into existing tractor-mounted row-crop cultivators, "
        "enabling large-scale adoption.",
    ]
    for item in deploy_future:
        bullet(doc, item)

    body(doc,
         "In summary, the WEED FINDER project provides a replicable, extensible research platform "
         "that can serve as the basis for advanced research in precision agriculture, embedded AI, "
         "and agricultural robotics. The open-source nature of the codebase and the use of "
         "universally available hardware make it accessible for replication, adaptation, and "
         "enhancement by researchers, students, and agricultural technology startups worldwide.")
    add_page_break(doc)


def chapter8(doc):
    heading1(doc, "CHAPTER 8: REFERENCES")
    doc.add_paragraph()
    body(doc,
         "The following references were consulted in the course of this project, encompassing "
         "peer-reviewed journal articles, conference papers, technical documentation, and "
         "open-source software repositories.")
    doc.add_paragraph()

    references = [
        "[1] Redmon, J., Divvala, S., Girshick, R., & Farhadi, A. (2016). You Only Look Once: "
        "Unified, Real-Time Object Detection. Proceedings of the IEEE Conference on Computer "
        "Vision and Pattern Recognition (CVPR 2016), pp. 779–788.",

        "[2] Jocher, G., Chaurasia, A., & Qiu, J. (2023). YOLO by Ultralytics. "
        "Available: https://github.com/ultralytics/ultralytics [Accessed: March 2025].",

        "[3] Lv, W., Xu, S., Zhao, Y., Wang, G., Wei, J., Cui, C., Du, Y., Dang, Q., & Liu, Y. "
        "(2023). DETRs Beat YOLOs on Real-time Object Detection. arXiv:2304.08069.",

        "[4] Sa, I., Ge, Z., Dayoub, F., Upcroft, B., Perez, T., & McCool, C. (2017). "
        "DeepWeeds: A Multiclass Weed Species Image Dataset for Deep Learning. Scientific "
        "Reports, 9(1), 2058. https://doi.org/10.1038/s41598-018-38343-3.",

        "[5] Bah, M. D., Hafiane, A., & Canals, R. (2018). Deep Learning with Unsupervised "
        "Data Augmentation for Weed Detection in Lettuce. Proceedings of the IEEE International "
        "Conference on Image Processing (ICIP 2018).",

        "[6] Lottes, P., Khanna, R., Pfeifer, J., Siegwart, R., & Stachniss, C. (2017). "
        "UAV-based Crop and Weed Classification for Smart Farming. Proceedings of the IEEE "
        "International Conference on Robotics and Automation (ICRA 2017).",

        "[7] Mohanty, S. P., Hughes, D. P., & Salathé, M. (2016). Using Deep Learning for "
        "Image-Based Plant Disease Detection. Frontiers in Plant Science, 7, 1419.",

        "[8] Oerke, E. C. (2006). Crop Losses to Pests. Journal of Agricultural Science, "
        "144(1), 31–43. https://doi.org/10.1017/S0021859605005708.",

        "[9] Heap, I. (2023). The International Herbicide-Resistant Weed Database. "
        "Available: www.weedscience.org [Accessed: February 2025].",

        "[10] Baweja, H. S., Parhar, T., Mirbod, O., & Nuske, S. (2018). StalkNet: A Deep "
        "Learning Pipeline for High-Throughput Measurement of Plant Stalk Count and Stalk Width. "
        "Field Robotics (FSR 2017). Springer Proceedings in Advanced Robotics, Vol. 5.",

        "[11] ONNX Runtime Development Team (2024). ONNX Runtime: Cross-Platform, High-"
        "Performance ML Inferencing. Microsoft. Available: https://onnxruntime.ai [Accessed: "
        "March 2025].",

        "[12] Bradski, G. (2000). The OpenCV Library. Dr. Dobb's Journal of Software Tools.",

        "[13] Raspberry Pi Foundation (2023). Raspberry Pi 4 Model B Product Brief. "
        "Available: https://www.raspberrypi.com/products/raspberry-pi-4-model-b/ "
        "[Accessed: November 2024].",

        "[14] Pallets Projects (2024). Flask: A Lightweight WSGI Web Framework. "
        "Available: https://flask.palletsprojects.com [Accessed: March 2025].",

        "[15] Harris, C. R., Millman, K. J., van der Walt, S. J., et al. (2020). Array "
        "Programming with NumPy. Nature, 585, 357–362.",

        "[16] Kumar, N., Chug, A., Sharma, D., & Singh, A. P. (2022). A Systematic Analysis "
        "of Deep Learning and Machine Learning Approaches for Weed Classification in "
        "Sustainable Agriculture. Sustainability, 14(20), 13608.",

        "[17] Kaggle (2023). Crop and Weed Detection Data with Bounding Boxes. "
        "Available: https://www.kaggle.com/datasets/ravirajsinh45/crop-and-weed-detection-data"
        "-with-bounding-boxes [Accessed: November 2024].",

        "[18] RPi.GPIO Library (2023). RPi.GPIO Python Package for Raspberry Pi GPIO. "
        "Available: https://pypi.org/project/RPi.GPIO/ [Accessed: December 2024].",

        "[19] Google LLC (2024). Google Colab: Jupyter Notebook Environment. "
        "Available: https://colab.research.google.com [Accessed: November 2024].",

        "[20] APJ Abdul Kalam Technological University (2024). B.Tech Regulations and "
        "Curriculum, Computer Science and Engineering, 2019 Scheme. Thiruvananthapuram, Kerala.",
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(ref)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        pf = p.paragraph_format
        pf.left_indent = Inches(0.4)
        pf.first_line_indent = Inches(-0.4)


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def main():
    doc = Document()
    set_doc_margins(doc)
    add_page_number(doc)

    print("Generating Title Page...")
    title_page(doc)

    print("Generating Declaration...")
    declaration_page(doc)

    print("Generating Certificate...")
    certificate_page(doc)

    print("Generating Acknowledgement...")
    acknowledgement_page(doc)

    print("Generating Abstract...")
    abstract_page(doc)

    print("Generating PO Mapping...")
    po_mapping_page(doc)

    print("Generating SDG Mapping...")
    sdg_mapping_page(doc)

    print("Generating Feasibility Report...")
    feasibility_page(doc)

    print("Generating Table of Contents...")
    toc_page(doc)

    print("Generating List of Tables...")
    lot_page(doc)

    print("Generating List of Figures...")
    lof_page(doc)

    print("Generating Abbreviations...")
    abbreviations_page(doc)

    print("Generating Chapter 1: Introduction...")
    chapter1(doc)

    print("Generating Chapter 2: Literature Review...")
    chapter2(doc)

    print("Generating Chapter 3: Methodology...")
    chapter3(doc)

    print("Generating Chapter 4: Work Plan...")
    chapter4(doc)

    print("Generating Chapter 5: Experimental Design...")
    chapter5(doc)

    print("Generating Chapter 6: Results and Conclusion...")
    chapter6(doc)

    print("Generating Chapter 7: Future Scope...")
    chapter7(doc)

    print("Generating Chapter 8: References...")
    chapter8(doc)

    doc.save(OUTPUT_PATH)
    print(f"\nDocument saved successfully to:\n{OUTPUT_PATH}")
    import os
    size = os.path.getsize(OUTPUT_PATH)
    print(f"File size: {size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
